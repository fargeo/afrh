# coding: utf-8
'''
ARCHES - a program developed to inventory and manage immovable cultural heritage.
Copyright (C) 2013 J. Paul Getty Trust and World Monuments Fund

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU Affero General Public License as
published by the Free Software Foundation, either version 3 of the
License, or (at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
GNU Affero General Public License for more details.

You should have received a copy of the GNU Affero General Public License
along with this program. If not, see <http://www.gnu.org/licenses/>.
'''

# import folium
import json
import os
import uuid
from datetime import datetime
import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Inches
# from io import BytesIO
from afrh_prj.utils.create_static_map import StaticMapCreator
from html.parser import HTMLParser
from html.entities import name2codepoint
from django.core.files.uploadedfile import UploadedFile
from django.http import HttpRequest, HttpResponseNotFound
from django.utils.translation import ugettext as _
from django.views.generic import View
from arches.app.datatypes.datatypes import DataTypeFactory
from arches.app.models import models
from arches.app.models.resource import Resource
from arches.app.models.system_settings import settings
from arches.app.models.tile import Tile
from arches.app.utils.betterJSONSerializer import JSONSerializer, JSONDeserializer
from arches.app.utils.response import JSONResponse
from arches.app.views.tile import TileData


class FileTemplateView(View):

    def __init__(self):
        self.doc = None
        self.resource = None
        self.date = None
        self.file_list_node_id = '2541f898-e0c7-11ea-8120-784f435179ea'

    def get(self, request):
        parenttile_id = request.GET.get('parenttile_id')
        parent_tile = Tile.objects.get(tileid=parenttile_id)
        letter_tiles = Tile.objects.filter(parenttile=parent_tile)
        url = None
        for tile in letter_tiles:
            if url is not None:
                break
            for data_obj in tile.data[self.file_list_node_id]:
                if data_obj['status'] == 'uploaded':
                    url = data_obj['url']
                    break

        if url is not None:
            return JSONResponse({'msg': 'success', 'download': url})
        return HttpResponseNotFound("No letters tile matching query by parent tile")
    
    def post(self, request): 
        datatype_factory = DataTypeFactory()
        parenttile_id = request.POST.get('parenttile_id')
        resourceinstance_id = request.POST.get('resourceinstance_id', None)
        self.resource = Resource.objects.get(resourceinstanceid=resourceinstance_id)
        self.resource.load_tiles()
        template_name = "afrh_1.docx"
        template_path = os.path.join(settings.APP_ROOT, 'docx', template_name)

        if not os.path.exists(os.path.join(settings.APP_ROOT, 'uploadedfiles','docx')):
            os.mkdir(os.path.join(settings.APP_ROOT, 'uploadedfiles','docx'))

        self.doc = Document(template_path)
        self.date = datetime.now().strftime("%Y-%m-%d")
        self.edit_letter_default(self.resource, datatype_factory)
        new_file_name = f"{self.date}_{template_name}"
        new_file_path = os.path.join(settings.APP_ROOT, 'uploadedfiles', 'docx', new_file_name)
        new_req = HttpRequest()
        new_req.method = 'POST'
        new_req.user = request.user
        new_req.POST['data'] = None
        host = request.get_host()

        self.doc.save(new_file_path)
        saved_file = open(new_file_path, 'rb')
        stat = os.stat(new_file_path)
        file_data = UploadedFile(saved_file)

        tile = json.dumps({
            "tileid":None,
            "data": {
                self.file_list_node_id: [{
                    "name":new_file_name,
                    "accepted":True,
                    "height":0,
                    "lastModified":stat.st_mtime,
                    "size":stat.st_size,
                    "status":"queued",
                    "type":"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "width":0,
                    "url":None,
                    "file_id":None,
                    "index":0,
                    "content":"blob:"+host+"/{0}".format(uuid.uuid4())
                }]
            },
            "nodegroup_id":self.file_list_node_id,
            "parenttile_id":parenttile_id,
            "resourceinstance_id":resourceinstance_id,
            "sortorder":0,
            "tiles":{}
        })

        new_req = HttpRequest()
        new_req.method = 'POST'
        new_req.user = request.user
        new_req.POST['data'] = tile
        new_req.FILES['file-list_' + self.file_list_node_id] = file_data
        new_tile_data_instance = TileData()
        post_resp = TileData.post(new_tile_data_instance, new_req)

        if post_resp.status_code == 200:
            return JSONResponse({'tile':tile, 'status':'success' })

        return HttpResponseNotFound(post_resp.status_code)

    def edit_letter_default(self, consultation, datatype_factory):
        template_dict = {
            'URR#':{'nodeid': '937529ba-3d6c-11ea-b9b7-027f24e6fd6b', 'default':'No URR# Entered', 'found': False},
            'Scope of Work Description': {'nodeid': '5a8422b0-3cac-11ea-b9b7-027f24e6fd6b', 'default':'No Scope of Work Description available', 'found': False},
            'Project Area Notes': {'nodeid': '9e69372e-779d-11ea-8977-acde48001122', 'default':'No Project Area Notes entered', 'found': False}, # aka Submission Notes
            # 'Character Areas List': '',
            # 'Master Plan Zones List': '',
            'Direct Impacts': {'nodeid': '344a48d8-f47a-11ea-a92a-a683e74f6c3a', 'default':'No Direct Impacts identified', 'found': False},
            'Indirect Impacts': {'nodeid': 'f36b5244-f479-11ea-a92a-a683e74f6c3a', 'default':'No Indirect Impacts identified', 'found': False},
            # 'APE Map': 'screenshot of this map',
            'AFRH Determination of Effect': {'nodeid': '7414718a-3d6b-11ea-b9b7-027f24e6fd6b', 'default':'No Determination of Effect identified', 'found': False}, # note graph spelling may differ ("Affect")
            'Submission Notes': {'nodeid': '9e69372e-779d-11ea-8977-acde48001122', 'default':'No Submission Notes entered', 'found': False},
            'AGENT': {'nodeid': 'b0007bfc-415e-11ea-b9b7-027f24e6fd6b', 'default':'No Agent identified', 'found': False},
            'AGENT TYPE': {'nodeid': '6da8cd54-3c8a-11ea-b9b7-027f24e6fd6b', 'default':'Agent Type unselected', 'found': False},
            'AFRH PROJECT CONTACT (Management Activity A, Entities)': {'nodeid': '6da8cd63-3c8a-11ea-b9b7-027f24e6fd6b', 'default':'No Project Contact identified', 'found': False},
            'Procedure Type (Management Activity A, Summary)': {'nodeid': 'feb5caf5-3c8b-11ea-b9b7-027f24e6fd6b', 'default':'No Procedure Type identified', 'found': False},
            'Documentation Type (Management Activity A, NEPA Review)': {'nodeid': '6da8cd45-3c8a-11ea-b9b7-027f24e6fd6b', 'default':'No Documentation Type identified', 'found': False},
        }
        self.replace_in_letter(consultation.tiles, template_dict, datatype_factory)

        direct_impact_nodegroupid = '344a48d8-f47a-11ea-a92a-a683e74f6c3a'
        indirect_impact_nodegroupid = 'f36b5244-f479-11ea-a92a-a683e74f6c3a'
        archeology_zone_graphid = 'ddb9385d-39fe-11ea-b9b7-027f24e6fd6b'
        master_plan_zone_graphid = '12581535-3a08-11ea-b9b7-027f24e6fd6b'
        character_area_graphid = 'f3ab0a3a-1aca-11ea-8211-acde48001122'
        activity_spatial_location_nodegroupid = '429130d2-6b27-11ea-b9b7-027f24e6fd6b'
        activity_spatial_location_coordinates_nodeid = '4d12497f-6b27-11ea-b9b7-027f24e6fd6b'
        related_arch_zone_resourceids = []
        master_plan_zones = []
        character_areas = []
        related_arch_zone_names = 'No Related Archaeology Zones'
        related_mpz_names = 'No Related Master Plan Zones'
        related_character_area_names = 'No Related Character Areas'
        within = False
        impact_tiles = list(filter(lambda x: (str(x.nodegroup_id) == direct_impact_nodegroupid or str(x.nodegroup_id) == indirect_impact_nodegroupid), consultation.tiles))

        # if one of the direct impacts is archaeology, grab the names of those resources from Direct Impacts, also set the checklist in doc
        for t in impact_tiles:
            if str(t.nodegroup_id) == direct_impact_nodegroupid:
                for related_res in t.data[direct_impact_nodegroupid]:
                    res = Resource.objects.get(pk=related_res["resourceId"])
                    if str(res.graph_id) == archeology_zone_graphid:
                        within = True
                        related_arch_zone_resourceids.append(related_res["resourceId"])

            elif str(t.nodegroup_id) == indirect_impact_nodegroupid:
                for related_res in t.data[indirect_impact_nodegroupid]:
                    res = Resource.objects.get(pk=related_res["resourceId"])
                    if str(res.graph_id) == master_plan_zone_graphid:
                        master_plan_zones.append(related_res["resourceId"])
                    elif str(res.graph_id) == character_area_graphid:
                        character_areas.append(related_res["resourceId"])

        if within:
            impact_dict = {
                'Within': '\u2612', # x box
                'Not within': '\u2610' # open box
            }
        else:
            impact_dict = {
                'Not within': '\u2612', # x box
                'Within': '\u2610' # open box
            }

        for r in Resource.objects.filter(pk__in=related_arch_zone_resourceids):
            if related_arch_zone_names == 'No Related Archaeology Zones':
                related_arch_zone_names = r.displayname
            else:
                related_arch_zone_names += (', ' + r.displayname)

        for r in Resource.objects.filter(pk__in=master_plan_zones):
            if related_mpz_names == 'No Related Master Plan Zones':
                related_mpz_names = r.displayname
            else:
                related_mpz_names += (', ' + r.displayname)

        for r in Resource.objects.filter(pk__in=character_areas):
            if related_character_area_names == 'No Related Character Areas':
                related_character_area_names = r.displayname
            else:
                related_character_area_names += (', ' + r.displayname)

        self.replace_string(self.doc, 'Related Archaeological Zones', related_arch_zone_names)
        self.replace_string(self.doc, 'Within', impact_dict['Within'])
        self.replace_string(self.doc, 'Not within', impact_dict['Not within'])
        self.replace_string(self.doc, 'Name (Master Plan Zones, Summary)', related_mpz_names)
        self.replace_string(self.doc, 'Name (Character Areas, Summary)', related_character_area_names)
        
        try:
            location = list(filter(lambda x: (str(x.nodegroup_id) == activity_spatial_location_nodegroupid), consultation.tiles))[0]
            map_output_path = os.path.join(settings.APP_ROOT, "docx", "temp", "temp_map.png")
            feature_collection = location.data[activity_spatial_location_coordinates_nodeid]
            static_map_creator = StaticMapCreator()
            static_map_creator.create_map(feature_collection, map_output_path, height=500, width=800)
            self.insert_image(self.doc, "APE Map", image_path=map_output_path)
        except (IndexError, KeyError):
            self.replace_string(self.doc, "APE Map", "\tNo APE Defined", usebraces=False)

    def replace_in_letter(self, tiles, template_dict, datatype_factory):
        self.replace_string(self.doc, 'AUTOMATIC DATE', self.date)
        for tile in tiles:
            for key, val_dict in list(template_dict.items()):
                html = False
                if val_dict['nodeid'] in tile.data: # nodeid is key in this tile.data
                    val_dict['found'] = True
                    my_node = models.Node.objects.get(nodeid=val_dict['nodeid'])
                    datatype = datatype_factory.get_instance(my_node.datatype)
                    lookup_val = datatype.get_display_value(tile, my_node)
                    if lookup_val is None or lookup_val == "":
                        lookup_val = val_dict['default']
                    if '<' in lookup_val: # not ideal for finding html/rtf
                        html = True
                    self.replace_string(self.doc, key, lookup_val, html)
                
        for key, val_dict in list(template_dict.items()): # for any fields remaining unpopulated, use default
            html = False
            if val_dict['found'] is False:
                lookup_val = val_dict['default']
                if '<' in lookup_val: # not ideal for finding html/rtf
                    html = True
                self.replace_string(self.doc, key, lookup_val, html)

    def replace_string(self, document, key, v, html=False, usebraces=True):
        # Note that the intent here is to preserve how things are styled in the docx
        # easiest way is to iterate through p.runs, not as fast as iterating through parent.paragraphs
        # advantage of the former is that replacing run.text preserves styling, replacing p.text does not
        
        def parse_html_to_docx(p, k, v):
            # style = p.style
            if k in p.text:
                p.clear()
                document_html_parser = DocumentHTMLParser(p, document)
                document_html_parser.insert_into_paragraph_and_feed(v)

        def replace_in_runs(p_list, k, v):
            for paragraph in p_list:
                if html is True:
                    parse_html_to_docx(paragraph, k, v)
                for i, run in enumerate(paragraph.runs):
                    if k in run.text:
                        run.text = run.text.replace(k, v)
                    elif i == (len(paragraph.runs) - 1) and k in paragraph.text:
                        paragraph.text = paragraph.text.replace(k, v)

        def iterate_tables(t_list, k, v):
            for table in t_list:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_runs(cell.paragraphs, k, v)
        
        if v is not None and key is not None:
            if usebraces:
                k = "{{"+key+"}}"
            else:
                k = key
            doc = document

            if len(doc.paragraphs) > 0:
                replace_in_runs(doc.paragraphs, k, v)

            if len(doc.tables) > 0:
                iterate_tables(doc.tables, k, v)
            
            if len(doc.sections) > 0:
                for section in doc.sections:
                    replace_in_runs(section.footer.paragraphs, k, v)
                    iterate_tables(section.footer.tables, k, v)
                    replace_in_runs(section.header.paragraphs, k, v)
                    iterate_tables(section.header.tables, k, v)

    
    def insert_image(self, document, k, image_path=None, image_obj=None, config=None):
        
        def replace_in_runs(p_list, k):
            for paragraph in p_list:
                for run in paragraph.runs:
                    if k in run.text and image_path:
                        run.text = run.text.replace(k, '')
                        run.add_picture(image_path)

        def iterate_tables(t_list, k):
            for table in t_list:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_runs(cell.paragraphs, k)
        
        if k is not None:
            if len(document.paragraphs) > 0:
                replace_in_runs(document.paragraphs, k)

        def insert_paragraph_after(self, paragraph, text=None, style=None):
            """Insert a new paragraph after the given paragraph."""
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            if style is not None:
                new_para.style = style
            return new_para

    def insert_custom(self, document, k, v, config=None):
        # perhaps replaces {{custom_object}} with pre-determined text structure with custom style/format
        return True

        
class DocumentHTMLParser(HTMLParser):
    def __init__(self, paragraph, document):
        HTMLParser.__init__(self)
        self.document = document
        self.paragraph = paragraph
        self.table = None
        self.table_cols = 0
        self.table_rows = 0
        self.max_cols_reached = False
        self.td_cursor = False
        self.hyperlink = False
        self.list_style = "ul"
        self.ol_counter = 1
        self.run = self.paragraph.add_run()

    def insert_paragraph_after(self, paragraph, text=None, style=None):
            """Insert a new paragraph after the given paragraph."""
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            if style is not None:
                new_para.style = style
            return new_para

    def add_hyperlink(self, paragraph, url, text, color=None, underline=None):
            # This gets access to the document.xml.rels file and gets a new relation id value
            part = self.paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

            # Create a w:r element
            new_run = docx.oxml.shared.OxmlElement('w:r')

            # Create a new w:rPr element
            rPr = docx.oxml.shared.OxmlElement('w:rPr')

            # Add color if it is given
            if not color is None:
                c = docx.oxml.shared.OxmlElement('w:color')
                c.set(docx.oxml.shared.qn('w:val'), color)
                rPr.append(c) # #5384da ; rgb(83,132,218)
            
            # Remove underlining if it is requested
            if not underline:
                u = docx.oxml.shared.OxmlElement('w:u')
                u.set(docx.oxml.shared.qn('w:val'), 'none')
                rPr.append(u)

            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

            return hyperlink

    def insert_into_paragraph_and_feed(self, html):
        self.run = self.paragraph.add_run()
        self.feed(html)

    def handle_starttag(self, tag, attrs):
        # print(tag,attrs)
        self.run = self.paragraph.add_run()
        if tag == "i" or tag == "em":
            self.run.italic = True
        if tag == "b" or tag == "strong":
            self.run.bold = True
        if tag == "s":
            self.run.strike = True
        if tag == "u":
            self.run.underline = True
        if tag == "ol":
            self.list_style = "ol"
        if tag == "ul":
            self.list_style = "ul"
        if tag in ["br", "ul", "ol"]:
            self.run.add_break()
        if tag == "li":
            if self.list_style == 'ul':
                self.run.add_text('● ')
            else:
                self.run.add_text(str(self.ol_counter)+'. ')
                self.ol_counter += 1
        if tag == "p":
            self.run.add_break()
            # self.run.add_break()
            # self.run.add_tab()
        if tag == "a":
            self.hyperlink = attrs[0][1]
        if tag == "table":
            self.table = self.document.add_table(self.table_rows, self.table_cols)
            self.table.autofit = True
        if tag == "tr":
            self.table_rows+= 1
            self.table.add_row()
        if tag == "td":
            self.table_cols+= 1
            if self.max_cols_reached is False:
                self.table.add_column(1)
            self.td_cursor = True

    def handle_endtag(self, tag):
        if tag in ["br", "li", "ul", "ol"]:
            self.run.add_break()
        self.run = self.paragraph.add_run()
        if tag == "ol":
            self.ol_counter = 1
        if tag == "table":
            tbl = self.table._tbl
            p = self.paragraph._p
            p.addnext(tbl)
            self.table = None
            self.table_cols = 0
            self.table_rows = 0
        if tag == "tr":
            self.table_cols = 0
            self.max_cols_reached = True
        if tag == "td":
            self.td_cursor = False

    def handle_data(self, data):
        if "&#39;" in data:
            data = data.replace("&#39;","\'")

        if self.hyperlink is not False:
            blue = docx.shared.RGBColor(83,132,218)
            color = blue.__str__()
            self.add_hyperlink(self.paragraph, self.hyperlink, data, color)
            self.hyperlink = False
        elif self.td_cursor is True:
            self.table.cell(self.table_rows-1, self.table_cols-1).add_paragraph(data)
        else:
            self.run.add_text(data)

    def handle_entityref(self, name):
        c = chr(name2codepoint[name])
        self.run.add_text(c)

    def handle_charref(self, name):
        if name.startswith('x'):
            c = chr(int(name[1:], 16))
        else:
            c = chr(int(name))
        self.run.add_text(c)

